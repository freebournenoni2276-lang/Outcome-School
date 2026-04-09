import pandas as pd
import os, sys
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

students = [
    {
        "name": "Aaliyah Brown",
        "grades": {"Math": 78, "English": 82, "Science": 75},
        "behaviour": "B",
        "attendance": "A",
        "effort": "B",
        "comment": "A reliable student who participates well and shows steady progress across subjects."
    },
    {
        "name": "Jaden Smith",
        "grades": {"Math": 65, "English": 70, "Science": 68},
        "behaviour": "C",
        "attendance": "B",
        "effort": "C",
        "comment": "Shows moderate understanding but needs more consistency in studying and class engagement."
    },
    {
        "name": "Maya Johnson",
        "grades": {"Math": 90, "English": 88, "Science": 92},
        "behaviour": "A",
        "attendance": "A",
        "effort": "A",
        "comment": "Outstanding student with excellent academic performance and strong work ethic."
    },
    {
        "name": "Daniel Williams",
        "grades": {"Math": 55, "English": 60, "Science": 58},
        "behaviour": "C",
        "attendance": "C",
        "effort": "D",
        "comment": "Struggles academically and requires additional support and improved study habits."
    },
    {
        "name": "Chloe Davis",
        "grades": {"Math": 84, "English": 79, "Science": 81},
        "behaviour": "B",
        "attendance": "A",
        "effort": "B",
        "comment": "Strong overall performance with good class participation and consistent effort."
    },
    {
        "name": "Ethan Wilson",
        "grades": {"Math": 72, "English": 68, "Science": 74},
        "behaviour": "B",
        "attendance": "B",
        "effort": "C",
        "comment": "Demonstrates decent understanding but needs to improve focus and revision habits."
    },
    {
        "name": "Zoe Taylor",
        "grades": {"Math": 95, "English": 91, "Science": 89},
        "behaviour": "A",
        "attendance": "A",
        "effort": "A",
        "comment": "Exceptional student with top-tier academic achievement and consistent excellence."
    },
    {
        "name": "Isaac Anderson",
        "grades": {"Math": 61, "English": 65, "Science": 63},
        "behaviour": "C",
        "attendance": "B",
        "effort": "C",
        "comment": "Shows potential but needs improvement in consistency and subject understanding."
    },
    {
        "name": "Amara Thomas",
        "grades": {"Math": 88, "English": 85, "Science": 87},
        "behaviour": "A",
        "attendance": "A",
        "effort": "B",
        "comment": "Excellent student with strong academic ability and positive classroom attitude."
    },
    {
        "name": "Noah Jackson",
        "grades": {"Math": 77, "English": 73, "Science": 70},
        "behaviour": "B",
        "attendance": "B",
        "effort": "B",
        "comment": "Good performance overall with steady progress and room for improvement in focus."
    }
]


# ---------- CALCULATIONS ----------

def calculate_average(grades):
    return round(sum(grades.values()) / len(grades), 2)


def letter_grade(grade):
    if grade >= 90:
        return "A"
    elif grade >= 80:
        return "B"
    elif grade >= 70:
        return "C"
    elif grade >= 60:
        return "D"
    else:
        return "F"


def student_gpa(letter):
    return {"A": 4.0, "B": 3.0, "C": 2.0, "D": 1.0}.get(letter, 0.0)


# Apply calculations
for student_data in students:
    student_data["average"] = calculate_average(student_data["grades"])

    student_data["letter_grades"] = {
        subject: letter_grade(score)
        for subject, score in student_data["grades"].items()
    }

    total = sum(student_gpa(l) for l in student_data["letter_grades"].values())
    student_data["GPA"] = round(total / len(student_data["letter_grades"]), 2)


# ---------- RANKING ----------

def class_ranking():
    sorted_students = sorted(students, key=lambda x: x["average"], reverse=True)
    for i, student in enumerate(sorted_students, start=1):
        student["position"] = i


# ---------------- ADDING STUDENTS -------------------------

def add_student():
    print("\n=== ADD NEW STUDENT ===")

    name = input("Enter student name: ").title()

    # ---------- MATH ----------
    while True:
        try:
            math = int(input("Enter Math score: "))
            if 0 <= math <= 100:
                break
            else:
                print("❌ Math score must be between 0 and 100")
        except ValueError:
            print("❌ Invalid input. Enter a number.")

    # ---------- ENGLISH ----------
    while True:
        try:
            english = int(input("Enter English score: "))
            if 0 <= english <= 100:
                break
            else:
                print("❌ English score must be between 0 and 100")
        except ValueError:
            print("❌ Invalid input. Enter a number.")

    # ---------- SCIENCE ----------
    while True:
        try:
            science = int(input("Enter Science score: "))
            if 0 <= science <= 100:
                break
            else:
                print("❌ Science score must be between 0 and 100")
        except ValueError:
            print("❌ Invalid input. Enter a number.")

    # ---------- PROCESS DATA ----------
    grades = {
        "Math": math,
        "English": english,
        "Science": science
    }

    average = calculate_average(grades)

    letter_grades = {
        subject: letter_grade(score)
        for subject, score in grades.items()
    }

    gpa_total = sum(student_gpa(l) for l in letter_grades.values())
    gpa = round(gpa_total / len(letter_grades), 2)

    # ---------- CONDUCT ----------
    def get_letter_input(prompt):
        while True:
            value = input(prompt).upper()
            if value in ["A", "B", "C", "D", "F"]:
                return value
            else:
                print("❌ Enter a valid grade (A, B, C, D, F)")

    behaviour = get_letter_input("Enter Behaviour (A-F): ")
    attendance = get_letter_input("Enter Attendance (A-F): ")
    effort = get_letter_input("Enter Effort (A-F): ")

     # ---------- COMMENT ----------
    comment = input("Enter teacher comment: ")


    new_student = {
        "name": name,
        "grades": grades,
        "average": average,
        "letter_grades": letter_grades,
        "GPA": gpa,
        "behaviour": behaviour,
        "attendance": attendance,
        "effort": effort,
        "comment": comment
    }

    students.append(new_student)

    print(f"\n✅ {name} added successfully!")


def export_student_word(student):
    doc = Document()

    # ================= PAGE SETUP (LEGAL + LANDSCAPE) =================
    section = doc.sections[0]

    section.page_width = Inches(14)   # Legal width (landscape)
    section.page_height = Inches(8.5)

    section.orientation = WD_ORIENT.LANDSCAPE

    # margins (tight to fit everything on one page)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # ================= HEADER =================
    title = doc.add_heading("STUDENT ACADEMIC TRANSCRIPT", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("Official Academic Report")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    # ================= STUDENT INFO =================
    doc.add_heading("Student Information", level=2)
    doc.add_paragraph(f"Name: {student['name']}   |   Position: {student.get('position', 'N/A')}")

    doc.add_paragraph("")

    # ================= TABLE (COMPACT) =================
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    hdr[0].text = "Subject"
    hdr[1].text = "Score"
    hdr[2].text = "Grade"

    for subject, score in student["grades"].items():
        row = table.add_row().cells
        row[0].text = subject
        row[1].text = str(score)
        row[2].text = student["letter_grades"][subject]

    doc.add_paragraph("")

    # ================= SUMMARY =================
    doc.add_heading("Summary", level=2)

    doc.add_paragraph(
        f"Average: {student['average']}   |   GPA: {student['GPA']}   |   "
        f"Status: {'PASS' if student['average'] >= 50 else 'FAIL'}"
    )

    doc.add_paragraph("")

    # ================= CONDUCT =================
    doc.add_heading("Conduct Evaluation", level=2)

    doc.add_paragraph(
        f"Behaviour: {student.get('behaviour', 'N/A')}   |   "
        f"Attendance: {student.get('attendance', 'N/A')}   |   "
        f"Effort: {student.get('effort', 'N/A')}"
    )

    doc.add_paragraph("")

    # ================= LEGEND =================
    doc.add_heading("Conduct Evaluation Key", level=2)
    doc.add_paragraph("A = Excellent   |   B = Good   |   C = Satisfactory   |   D = Needs Improvement   |   F = Poor")
    
    doc.add_paragraph("")

    # ================= COMMENT =================
    doc.add_heading("Teacher Comment", level=2)
    doc.add_paragraph(student.get("comment", "No comment provided."))


    # ================= FOOTER =================
    doc.add_paragraph("")
    footer = doc.add_paragraph("Official School Record - Generated by Student Management System")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ================= SAVE =================
    filename = f"{student['name'].replace(' ', '_')}_Transcript.docx"
    doc.save(filename)

    print(f"\n📄 One-page transcript saved as {filename}")

    # ================= AUTO OPEN =================
    try:
        if sys.platform == "win32":
            os.startfile(filename)
    except OSError as e:
        print("Could not auto-open file:", e)

 # ================= STUDENT AVERAGE CHART =================
def plot_student_averages():
    names = [s["name"] for s in students]
    averages = [s["average"] for s in students]

    plt.figure()
    plt.bar(names, averages)
    plt.xticks(rotation=45, ha="right")
    plt.title("Student Average Scores")
    plt.ylabel("Average Score")
    plt.tight_layout()
    plt.show()

# ================= STUDENT PERFORMANCE CHART =================
def plot_subject_averages():
    subjects = ["Math", "English", "Science"]

    averages = []
    for subject in subjects:
        avg = sum(s["grades"][subject] for s in students) / len(students)
        averages.append(avg)

    plt.figure()
    plt.bar(subjects, averages)
    plt.title("Class Subject Performance")
    plt.ylabel("Average Score")
    plt.show()

# ================= CLASS SUMMARY DASHBOARD =================
def class_summary():
    avg_class = sum(s["average"] for s in students) / len(students)

    top = max(students, key=lambda x: x["average"])
    low = min(students, key=lambda x: x["average"])

    pass_count = sum(1 for s in students if s["average"] >= 50)

    print("\n===== CLASS DASHBOARD =====")
    print(f"Class Average: {round(avg_class,2)}")
    print(f"Top Student: {top['name']} ({top['average']})")
    print(f"Lowest Student: {low['name']} ({low['average']})")
    print(f"Pass Rate: {pass_count}/{len(students)}")

# ================= EXPORT ALL REPORTS =================
def export_all_word_reports():
    class_ranking()

    for student in students:
        export_student_word(student)

    print("\n📄 All student transcripts exported successfully!")


# ---------- SEARCH ----------

def search_function():
    search = input("Enter student name: ").strip().lower()

    for student in students:
        if student["name"].lower() == search:

            print("\n==============================")
            print("      STUDENT REPORT")
            print("==============================")

            print(f"\nName: {student['name']}")
            print(f"Position: {student.get('position', 'Not ranked')}")

            print("\nSubjects:")
            for subject, score in student["grades"].items():
                letter = student["letter_grades"][subject]
                print(f"{subject:<10}: {score} ({letter})")

            # ✅ OUTSIDE the loop → prints ONCE
            print("\n------------------------------")
            print(f"Average : {student['average']}")
            print(f"GPA     : {student['GPA']}")

            status = "Pass" if student["average"] >= 50 else "Fail"
            print(f"Status  : {status}")

            print("==============================\n")

            choice = input("Export this report to Word? (y/n): ").lower()

            if choice == "y":
                export_student_word(student)

            return

    print("Student Not Found")


# ---------- EXPORT ----------

def export_reports():
    class_ranking()

    data = []
    for student in students:
        data.append({
            "Name": student["name"],
            "Math": student["grades"]["Math"],
            "English": student["grades"]["English"],
            "Science": student["grades"]["Science"],
            "Average": student["average"],
            "GPA": student["GPA"],
            "Position": student["position"]
        })

    df = pd.DataFrame(data).sort_values(by="Position")

    csv_file = "Student_Report.csv"
    excel_file = "Student_Report.xlsx"

    df.to_csv(csv_file, index=False)
    try:
        df.to_excel(excel_file, index=False)
    except PermissionError:
        print("⚠️ Close the Excel file before exporting.")

    print("Reports exported successfully!")

    # 🔥 AUTO OPEN FILES
    os.startfile(csv_file)
    os.startfile(excel_file)

# ---------- MAIN MENU ----------

def main_menu():
    while True:
        print("\n============ EDUTRACK ===============")
        print("\n===== STUDENT MANAGEMENT SYSTEM =====")
        print("1. Search Student Report")
        print("2. Add New Student")
        print("3. View Class Summary Dashboard")
        print("4. Student Average Chart")
        print("5. Subject Performance Chart")
        print("6. Export Reports (CSV & Excel)")
        print("7. Export ALL Word Reports")
        print("8. Exit")

        choice = input("Choose an option: ")

        if choice == "1":
            class_ranking()
            search_function()
        
        elif choice == "2":
            add_student()

        elif choice == "3":
            class_ranking()
            class_summary()

        elif choice == "4":
            plot_student_averages()

        elif choice == "5":
            plot_subject_averages()

        elif choice == "6":
            export_reports()

        elif choice == "7":
            export_all_word_reports()

        elif choice == "8":
            print("Goodbye 👋")
            break

        else:
            print("Invalid choice, try again.")


main_menu()

